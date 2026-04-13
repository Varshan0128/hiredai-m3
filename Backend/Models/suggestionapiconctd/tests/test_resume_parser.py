from pathlib import Path

from docx import Document

from services.resume_parser import ResumeParserService


def test_docx_resume_parser_extracts_text(tmp_path: Path) -> None:
    file_path = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("Python Developer")
    document.add_paragraph("FastAPI and PostgreSQL")
    document.save(file_path)

    parser = ResumeParserService()

    text = parser.extract_text(file_path)

    assert "Python Developer" in text
    assert "FastAPI" in text


def test_pdf_resume_parser_uses_pdf_reader(monkeypatch, tmp_path: Path) -> None:
    file_path = tmp_path / "resume.pdf"
    file_path.write_bytes(b"%PDF-1.4 fake")

    class FakePage:
        @staticmethod
        def extract_text() -> str:
            return "Backend Developer with Python"

    class FakeReader:
        def __init__(self, _handle) -> None:
            self.pages = [FakePage()]

    monkeypatch.setattr("services.resume_parser.PyPDF2.PdfReader", FakeReader)

    parser = ResumeParserService()

    text = parser.extract_text(file_path)

    assert "Backend Developer" in text
