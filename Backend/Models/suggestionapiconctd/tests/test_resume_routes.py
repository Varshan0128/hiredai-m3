import os
from pathlib import Path
from uuid import uuid4

from docx import Document
from fastapi.testclient import TestClient


os.environ.setdefault("DATABASE_URL", "sqlite:///./test_resume_profiles.db")

from app import app  # noqa: E402


def _build_resume_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("Asha Kumar")
    document.add_paragraph("Career Objective")
    document.add_paragraph("Aspiring Project Manager seeking Project Coordinator opportunities.")
    document.add_paragraph("Certifications")
    document.add_paragraph("Google Project Management Certification")
    document.add_paragraph("Skills")
    document.add_paragraph("Project Management, Leadership, Communication, Coordination")
    document.add_paragraph("Projects")
    document.add_paragraph("Coordinated a no-code AI prototype presentation and handled stakeholder communication.")
    document.add_paragraph("Languages: English, Tamil")
    document.save(path)


def test_upload_and_fetch_resume_profile(tmp_path: Path) -> None:
    resume_path = tmp_path / "resume.docx"
    _build_resume_docx(resume_path)
    user_id = str(uuid4())

    with TestClient(app) as client:
        with resume_path.open("rb") as handle:
            upload_response = client.post(
                "/api/v1/resume/upload",
                data={"user_id": user_id},
                files={"file": ("resume.docx", handle, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )

        assert upload_response.status_code == 201, upload_response.text
        payload = upload_response.json()
        assert payload["profile"]["user_id"] == user_id
        assert "Project Management" in [item["name"] for item in payload["profile"]["explicit_skills"]]
        assert "Project Coordinator" in [item["title"] for item in payload["profile"]["recommended_roles"]]
        assert "Git" not in [item["name"] for item in payload["profile"]["explicit_skills"]]

        profile_response = client.get(f"/api/v1/resume/profile/{user_id}")
        assert profile_response.status_code == 200, profile_response.text
        assert profile_response.json()["original_filename"] == "resume.docx"


def test_upload_rejects_invalid_file_type() -> None:
    user_id = str(uuid4())
    with TestClient(app) as client:
        response = client.post(
            "/api/v1/resume/upload",
            data={"user_id": user_id},
            files={"file": ("resume.txt", b"plain text", "text/plain")},
        )

        assert response.status_code == 400
        assert "Only PDF and DOCX" in response.json()["detail"]
